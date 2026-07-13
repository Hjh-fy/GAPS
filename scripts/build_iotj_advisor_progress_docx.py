"""Build the advisor-facing IoT-J progress brief and reusable result figures."""
from __future__ import annotations

import argparse
import csv
import re
from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = (
    REPO_ROOT / "docs" / "paper" / "GAPS_IoTJ_advisor_progress_report_20260713.zh.md"
)
DEFAULT_OUTPUT = (
    REPO_ROOT / "docs" / "paper" / "GAPS_IoTJ_advisor_progress_report_20260713.zh.docx"
)
ASSET_DIR = REPO_ROOT / "docs" / "paper" / "assets" / "advisor_progress_20260713"
A_METRICS = (
    REPO_ROOT
    / "results"
    / "iotj_classification_ablation_20260711_v2r1_summary"
    / "classification_per_run.csv"
)
B_METRICS = (
    REPO_ROOT
    / "results"
    / "iotj_classification_ablation_20260712_v3_summary"
    / "classification_per_run.csv"
)
REGRESSION_METRICS = (
    REPO_ROOT
    / "results"
    / "iotj_c5_formal_regression_20260713_v2_summary"
    / "r0_r7_comparison.csv"
)
QC_METRICS = (
    REPO_ROOT
    / "results"
    / "iotj_c5_formal_regression_20260713_v2_summary"
    / "qc_operational_comparison.csv"
)
MSYH_FONT = Path("C:/Windows/Fonts/msyh.ttc")

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
HEADER_FOOTER_DXA = 708

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "202832"
MUTED = "66717D"
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
LIGHT_BLUE = "DCEBF7"
LIGHT_GREEN = "DCEFE4"
LIGHT_GOLD = "F8EBC7"
LIGHT_CORAL = "F8E2D7"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(
    run,
    *,
    ascii_font: str = "Calibri",
    east_asia_font: str = "Microsoft YaHei",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: float, color: str, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_paragraph_border(paragraph, *, color: str, size: int = 10) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_callout_box(paragraph, *, fill: str, border: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:space"), "5")
        node.set(qn("w:color"), border)
        p_bdr.append(node)
    p_pr.append(p_bdr)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, display, end))
    set_run_font(run, size=8.5, color=MUTED)


def configure_header_footer(section) -> None:
    section.header_distance = Twips(HEADER_FOOTER_DXA)
    section.footer_distance = Twips(HEADER_FOOTER_DXA)
    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header.add_run("GAPS 云边协同气体感知系统")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    header.add_run("\t")
    right = header.add_run("导师进展汇报 | 2026-07-14")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    prefix = footer.add_run("GAPS IoT-J Progress Brief  |  ")
    set_run_font(prefix, size=8.5, color=MUTED)
    add_page_field(footer)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, size=13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    for name, size, color, italic, align in (
        ("Brief Caption", 9, MUTED, True, WD_ALIGN_PARAGRAPH.CENTER),
        ("Table Citation", 9, MUTED, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Brief Equation", 10.5, DARK_BLUE, False, WD_ALIGN_PARAGRAPH.CENTER),
    ):
        style = doc.styles[name] if name in doc.styles else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, size=size, color=color)
        style.font.italic = italic
        style.paragraph_format.alignment = align
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(8 if name != "Table Citation" else 4)
        style.paragraph_format.keep_together = True


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        configure_header_footer(section)
    configure_styles(doc)
    doc.core_properties.title = "GAPS 云边协同气体感知系统论文进展汇报"
    doc.core_properties.subject = "IoT-J advisor progress brief"
    doc.core_properties.author = "GAPS Research Team"
    doc.core_properties.keywords = "cloud-edge; gas sensing; calibration; regression; quality control"


def next_numbering_ids(numbering) -> tuple[int, int]:
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    return max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1


def create_numbering(doc: Document, *, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id, num_id = next_numbering_ids(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend((level, number))
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


INLINE_PATTERN = re.compile(r"(\*\*.*?\*\*|`.*?`)")


def add_inline_text(paragraph, text: str, *, size: float | None = None) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                ascii_font="Consolas",
                east_asia_font="Microsoft YaHei",
                size=9.3 if size is None else min(size, 9.3),
                color=DARK_BLUE,
            )
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size)


def add_callout(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Twips(TABLE_INDENT_DXA)
    paragraph.paragraph_format.right_indent = Twips(TABLE_INDENT_DXA)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.keep_together = True
    set_callout_box(paragraph, fill=CALLOUT_FILL, border="8A98A5")
    label = paragraph.add_run("汇报结论  ")
    set_run_font(label, size=10, color=BLUE, bold=True)
    add_inline_text(paragraph, text, size=10)


def add_masthead(doc: Document, title: str) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(2)
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("ADVISOR PROGRESS BRIEF  |  IoT-J SYSTEM PAPER")
    set_run_font(run, size=9, color=BLUE, bold=True)

    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(6)
    title_paragraph.paragraph_format.line_spacing = 1.05
    title_paragraph.paragraph_format.keep_with_next = True
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, size=22, color="000000", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("分类消融、个性化浓度回归、可靠性 QC 与论文闭环")
    set_run_font(subtitle_run, size=13, color=DARK_BLUE, bold=True)

    metadata = (
        ("汇报对象", "导师"),
        ("主协议", "C1/C2 source -> C5 target"),
        ("实验平台", "Alibaba Cloud ECS + physical Raspberry Pi + Windows PC"),
        ("证据状态", "seed-42 主闭环完成；跨方向与配对多种子进行中"),
        ("汇报日期", "2026-07-14"),
    )
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=9.5, color=INK, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=9.5, color=INK)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(3)
    rule.paragraph_format.space_after = Pt(10)
    set_paragraph_border(rule, color=BLUE, size=14)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def setup_matplotlib():
    import matplotlib
    import matplotlib.pyplot as plt
    from matplotlib import font_manager

    matplotlib.use("Agg")
    if MSYH_FONT.is_file():
        font_manager.fontManager.addfont(str(MSYH_FONT))
        font_name = font_manager.FontProperties(fname=str(MSYH_FONT)).get_name()
        plt.rcParams["font.sans-serif"] = [font_name]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["font.size"] = 10
    return plt


def save_figure(fig, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    fig.clf()


def build_architecture_figure(path: Path) -> None:
    plt = setup_matplotlib()
    from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

    fig, ax = plt.subplots(figsize=(12, 5.2))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5.2)
    ax.axis("off")

    def box(x, y, w, h, title, detail, color):
        patch = FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.025,rounding_size=0.06",
            linewidth=1.2, edgecolor="#607080", facecolor=color,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h * 0.65, title, ha="center", va="center", fontsize=11, fontweight="bold", color="#1F2A35")
        ax.text(x + w / 2, y + h * 0.30, detail, ha="center", va="center", fontsize=8.6, color="#44515E", linespacing=1.25)

    def arrow(start, end, label=""):
        ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=13, linewidth=1.4, color="#526474"))
        if label:
            ax.text((start[0] + end[0]) / 2, (start[1] + end[1]) / 2 + 0.16, label, ha="center", va="bottom", fontsize=8.2, color="#526474")

    box(0.2, 3.25, 2.1, 1.15, "树莓派 C1", "源域窗口\n本地分类训练", "#DCEBF7")
    box(0.2, 1.05, 2.1, 1.15, "Windows PC C2", "源域窗口\n本地分类训练", "#DCEBF7")
    box(3.0, 2.15, 2.35, 1.25, "阿里云 ECS", "Flower 聚合\n语义记忆与服务器适应", "#DCEFE4")
    box(3.0, 0.25, 2.35, 1.1, "C5 Calibration", "320 窗口\n类别与浓度标签", "#F8EBC7")
    box(6.05, 2.15, 2.1, 1.25, "跨域分类路由", "预测气体类别\n置信度 + 64 维特征", "#E8EEF5")
    box(8.75, 2.15, 2.2, 1.25, "C5 个性化回归", "Ridge / MLP 候选\n当前部署最优为 R4 Ridge", "#F8E2D7")
    box(8.75, 0.25, 2.2, 1.1, "高覆盖率 QC", "accept / review / reject\n部署可见风险", "#E8EEF5")
    box(6.05, 0.25, 2.1, 1.1, "可靠输出", "自动 ppm\n或转入复核", "#DCEFE4")
    arrow((2.3, 3.8), (3.0, 2.95), "参数/统计")
    arrow((2.3, 1.6), (3.0, 2.55), "参数/统计")
    arrow((4.18, 1.35), (4.18, 2.15), "校准辅助")
    arrow((5.35, 2.78), (6.05, 2.78), "冻结分类器")
    arrow((8.15, 2.78), (8.75, 2.78), "预测路由")
    arrow((9.85, 2.15), (9.85, 1.35), "风险评分")
    arrow((8.75, 0.8), (8.15, 0.8), "自动/复核")
    ax.text(6.0, 4.75, "真实云边训练 + 目标校准 + 个性化回归 + 可靠输出", ha="center", va="center", fontsize=14, fontweight="bold", color="#1F4D78")
    save_figure(fig, path)
    plt.close(fig)


def build_classification_a_figure(path: Path) -> None:
    plt = setup_matplotlib()
    rows = {row["group_id"]: row for row in read_csv(A_METRICS)}
    groups = ["A0", "A0T", "A2", "A3", "A4", "A4S", "A5", "A6", "A7"]
    values = [100 * float(rows[group]["accuracy"]) for group in groups]
    colors = ["#8795A1", "#2A9D8F", "#AAB4BE", "#AAB4BE", "#AAB4BE", "#AAB4BE", "#E9C46A", "#4C956C", "#2E74B5"]
    fig, ax = plt.subplots(figsize=(10.5, 4.7))
    bars = ax.bar(groups, values, color=colors, edgecolor="#FFFFFF", linewidth=0.8)
    ax.set_ylim(0, 105)
    ax.set_ylabel("C5 Test Accuracy (%)")
    ax.set_title("第一层消融：服务器校准适应是跨设备分类的关键", fontweight="bold", color="#1F4D78")
    ax.grid(axis="y", alpha=0.22)
    ax.spines[["top", "right"]].set_visible(False)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 1.3, f"{value:.2f}", ha="center", va="bottom", fontsize=8.5)
    fig.tight_layout()
    save_figure(fig, path)
    plt.close(fig)


def build_classification_b_figure(path: Path) -> None:
    plt = setup_matplotlib()
    rows = {row["group_id"]: row for row in read_csv(B_METRICS)}
    groups = ["B1", "B2", "B3", "B4", "B5"]
    values = [100 * float(rows[group]["accuracy"]) for group in groups]
    errors = [1360 - round(1360 * float(rows[group]["accuracy"])) for group in groups]
    colors = ["#6C8EBF", "#2A9D8F", "#8C9FB1", "#5F7F8F", "#D97B66"]
    fig, ax = plt.subplots(figsize=(9.8, 4.7))
    bars = ax.bar(groups, values, color=colors, edgecolor="#FFFFFF", linewidth=0.8)
    ax.set_ylim(98.4, 99.45)
    ax.set_ylabel("C5 Test Accuracy (%)")
    ax.set_title("修正后 B1-B5：完整模块叠加没有超过紧凑 B2", fontweight="bold", color="#1F4D78")
    ax.grid(axis="y", alpha=0.22)
    ax.spines[["top", "right"]].set_visible(False)
    for bar, value, error in zip(bars, values, errors):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 0.025, f"{value:.3f}%\n{error} errors", ha="center", va="bottom", fontsize=8.7)
    fig.tight_layout()
    save_figure(fig, path)
    plt.close(fig)


def build_regression_figure(path: Path) -> None:
    plt = setup_matplotlib()
    rows = read_csv(REGRESSION_METRICS)
    classifiers = ["A6", "B5", "B2"]
    all_values = []
    correct_values = []
    for classifier in classifiers:
        selected = [
            row for row in rows
            if row["classifier_id"] == classifier and row["mode"] == "R4"
        ]
        by_scope = {row["scope"]: row for row in selected}
        all_values.append(float(by_scope["S_ALL"]["RMSE"]))
        correct_values.append(float(by_scope["S_CC"]["RMSE"]))
    import numpy as np

    x = np.arange(len(classifiers))
    width = 0.34
    fig, ax = plt.subplots(figsize=(9.6, 4.9))
    bars_all = ax.bar(x - width / 2, all_values, width, label="S_ALL 真实路由", color="#D97B66")
    bars_cc = ax.bar(x + width / 2, correct_values, width, label="S_CC 分类正确", color="#2A9D8F")
    ax.set_xticks(x, classifiers)
    ax.set_ylabel("R4 RMSE (ppm)")
    ax.set_ylim(0, 32)
    ax.set_title("端到端回归差异主要来自分类错路由", fontweight="bold", color="#1F4D78")
    ax.legend(frameon=False, ncol=2, loc="upper right")
    ax.grid(axis="y", alpha=0.22)
    ax.spines[["top", "right"]].set_visible(False)
    for bars in (bars_all, bars_cc):
        for bar in bars:
            value = bar.get_height()
            ax.text(bar.get_x() + bar.get_width() / 2, value + 0.55, f"{value:.2f}", ha="center", va="bottom", fontsize=9)
    fig.tight_layout()
    save_figure(fig, path)
    plt.close(fig)


def build_qc_figure(path: Path) -> None:
    plt = setup_matplotlib()
    rows = [row for row in read_csv(QC_METRICS) if row["classifier_id"] == "B2"]
    by_workpoint = {row["workpoint"]: row for row in rows}
    workpoints = ["FULL", "HC95", "HC90"]
    yield_values = [100 * float(by_workpoint[item]["automatic_yield"]) for item in workpoints]
    rmse_values = [float(by_workpoint[item]["accept_RMSE"]) for item in workpoints]
    route_recall = [100 * float(by_workpoint[item]["route_wrong_recall"]) for item in workpoints]
    fig, ax = plt.subplots(figsize=(9.6, 4.9))
    ax.plot(yield_values, rmse_values, color="#2E74B5", linewidth=2.2, marker="o", markersize=8)
    colors = ["#8795A1", "#2A9D8F", "#E9C46A"]
    for workpoint, x_value, y_value, recall, color in zip(workpoints, yield_values, rmse_values, route_recall, colors):
        ax.scatter([x_value], [y_value], s=90, color=color, edgecolor="white", linewidth=1.0, zorder=3)
        ax.annotate(
            f"{workpoint}\nRMSE {y_value:.2f}\n错路由召回 {recall:.0f}%",
            (x_value, y_value), xytext=(0, 12), textcoords="offset points",
            ha="center", va="bottom", fontsize=8.7,
        )
    ax.set_xlim(87.5, 101.5)
    ax.set_ylim(10.7, 15.6)
    ax.set_xlabel("自动输出率 (%)")
    ax.set_ylabel("Accepted RMSE (ppm)")
    ax.set_title("B2 高覆盖率 QC：少量分流换取集中风险发现", fontweight="bold", color="#1F4D78")
    ax.grid(alpha=0.22)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    save_figure(fig, path)
    plt.close(fig)


def generate_figures() -> dict[str, Path]:
    required = (A_METRICS, B_METRICS, REGRESSION_METRICS, QC_METRICS)
    missing = [path for path in required if not path.is_file()]
    if missing:
        raise FileNotFoundError(f"missing evidence files: {missing}")
    figures = {
        "system_architecture.png": ASSET_DIR / "system_architecture.png",
        "classification_a_groups.png": ASSET_DIR / "classification_a_groups.png",
        "classification_b_groups.png": ASSET_DIR / "classification_b_groups.png",
        "regression_r4_comparison.png": ASSET_DIR / "regression_r4_comparison.png",
        "qc_tradeoff.png": ASSET_DIR / "qc_tradeoff.png",
    }
    build_architecture_figure(figures["system_architecture.png"])
    build_classification_a_figure(figures["classification_a_groups.png"])
    build_classification_b_figure(figures["classification_b_groups.png"])
    build_regression_figure(figures["regression_r4_comparison.png"])
    build_qc_figure(figures["qc_tradeoff.png"])
    return figures


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(cells: Sequence[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


TABLE_WIDTHS: dict[tuple[str, ...], tuple[int, ...]] = {
    ("项目", "设置"): (1800, 7560),
    ("配置", "共同核心", "额外分布项"): (1100, 3650, 4610),
    ("组别", "C5 Accuracy", "作用解释"): (850, 1500, 7010),
    ("组别", "新增项", "Accuracy", "Macro-F1", "错误数"): (700, 3020, 1650, 1650, 2340),
    ("编号", "模型/策略", "作用与结论"): (650, 2350, 6360),
    ("分类骨干", "分类错误数", "R4 S_ALL RMSE", "R4 S_CC N", "R4 S_CC RMSE"): (1200, 1450, 2250, 1600, 2860),
    ("工作点", "Accept/Review/Reject", "自动输出率", "Nonreject 覆盖率", "Accepted RMSE", "错路由召回"): (900, 2050, 1250, 1450, 1700, 2010),
    ("证据线", "当前最好可部署结果", "结论等级"): (1450, 5100, 2810),
}


def choose_table_widths(headers: Sequence[str]) -> tuple[int, ...]:
    key = tuple(headers)
    if key in TABLE_WIDTHS:
        return TABLE_WIDTHS[key]
    count = len(headers)
    base = CONTENT_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return tuple(widths)


def add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    headers = list(rows[0])
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(value)
        set_run_font(run, size=8.3, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_values in rows[1:]:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if index > 0 and re.fullmatch(r"[-+]?\d+(?:\.\d+)?%?(?: ppm)?", value)
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            add_inline_text(paragraph, value, size=8.3)
    set_table_geometry(table, choose_table_widths(headers))
    spacer = doc.add_paragraph(style="Table Citation")
    spacer.add_run("")


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(5.9))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    cap = doc.add_paragraph(style="Brief Caption")
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=9, color=MUTED, italic=True)


def parse_markdown(doc: Document, source: Path, figures: dict[str, Path]) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    title = next((line[2:].strip() for line in lines if line.startswith("# ")), "GAPS 论文进展汇报")
    add_masthead(doc, title)
    index = 1
    current_section = ""
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("**用途：") or line.startswith("**主协议：") or line.startswith("**实验拓扑：") or line.startswith("**结果截点：") or line.startswith("**建议汇报时长："):
            index += 1
            continue
        if line.startswith(">"):
            callout = line.lstrip("> ")
            add_callout(doc, callout)
            index += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            index += 1
            continue
        if line.startswith("## "):
            current_section = line[3:].strip()
            doc.add_paragraph(current_section, style="Heading 1")
            index += 1
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            caption, relative = image_match.groups()
            image_path = (source.parent / relative).resolve()
            if not image_path.is_file():
                candidate = figures.get(Path(relative).name)
                if candidate is None or not candidate.is_file():
                    raise FileNotFoundError(image_path)
                image_path = candidate
            add_figure(doc, image_path, caption)
            index += 1
            continue
        if line == r"\[":
            equation_lines: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != r"\]":
                equation_lines.append(lines[index].strip())
                index += 1
            paragraph = doc.add_paragraph(style="Brief Equation")
            run = paragraph.add_run(" ".join(equation_lines))
            set_run_font(run, ascii_font="Cambria Math", size=10.5, color=DARK_BLUE)
            index += 1
            continue
        if line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            parsed = [parse_table_row(item) for item in table_lines]
            parsed = [row for row in parsed if not is_separator_row(row)]
            add_table(doc, parsed)
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            items: list[str] = []
            while index < len(lines):
                match = re.match(r"^\d+\.\s+(.*)$", lines[index].strip())
                if not match:
                    break
                items.append(match.group(1))
                index += 1
            num_id = create_numbering(doc, kind="decimal")
            for item in items:
                paragraph = doc.add_paragraph(style="Normal")
                apply_numbering(paragraph, num_id)
                if current_section == "六、证据文件索引":
                    paragraph.paragraph_format.space_after = Pt(3)
                    paragraph.paragraph_format.line_spacing = 1.05
                    add_inline_text(paragraph, item, size=9)
                else:
                    add_inline_text(paragraph, item)
            continue
        bullet = re.match(r"^-\s+(.*)$", line)
        if bullet:
            items = []
            while index < len(lines):
                match = re.match(r"^-\s+(.*)$", lines[index].strip())
                if not match:
                    break
                items.append(match.group(1))
                index += 1
            num_id = create_numbering(doc, kind="bullet")
            for item in items:
                paragraph = doc.add_paragraph(style="Normal")
                apply_numbering(paragraph, num_id)
                add_inline_text(paragraph, item)
            continue
        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line:
                break
            if (
                next_line.startswith(("## ", "### ", "|", ">", "![", r"\["))
                or re.match(r"^\d+\.\s+", next_line)
                or re.match(r"^-\s+", next_line)
            ):
                break
            paragraph_lines.append(next_line)
            index += 1
        paragraph = doc.add_paragraph(style="Normal")
        if len(paragraph_lines) == 1 and re.fullmatch(r"\*\*[^*]+\*\*", line):
            paragraph.paragraph_format.keep_with_next = True
        add_inline_text(paragraph, " ".join(paragraph_lines))


def audit_document(doc: Document) -> None:
    section = doc.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.top_margin == Inches(1)
    assert section.right_margin == Inches(1)
    assert section.bottom_margin == Inches(1)
    assert section.left_margin == Inches(1)
    assert section.header_distance == Twips(HEADER_FOOTER_DXA)
    assert section.footer_distance == Twips(HEADER_FOOTER_DXA)
    normal = doc.styles["Normal"]
    assert normal.font.size == Pt(11)
    assert normal.paragraph_format.space_after == Pt(6)
    assert abs(float(normal.paragraph_format.line_spacing) - 1.10) < 0.005
    for name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = doc.styles[name]
        assert style.font.size == Pt(size)
        assert style.paragraph_format.space_before == Pt(before)
        assert style.paragraph_format.space_after == Pt(after)
    if len(doc.inline_shapes) != 5:
        raise AssertionError(f"expected five figures, found {len(doc.inline_shapes)}")
    if not doc.tables:
        raise AssertionError("brief must contain result tables")
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        layout = tbl_pr.find(qn("w:tblLayout"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == CONTENT_WIDTH_DXA
        assert tbl_ind is not None and int(tbl_ind.get(qn("w:w"))) == TABLE_INDENT_DXA
        assert layout is not None and layout.get(qn("w:type")) == "fixed"
        grid_widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
        assert sum(grid_widths) == CONTENT_WIDTH_DXA
    body = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    for required in ("明天建议重点汇报的四个结果", "修正后 B1-B5", "QC 如何在高覆盖率下工作", "导师可能追问"):
        if required not in body:
            raise AssertionError(f"missing required content: {required}")


def build_document(source: Path, output: Path) -> Path:
    if not source.is_file():
        raise FileNotFoundError(source)
    figures = generate_figures()
    doc = Document()
    configure_document(doc)
    parse_markdown(doc, source, figures)
    audit_document(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    reloaded = Document(output)
    audit_document(reloaded)
    return output


def main(argv: Sequence[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args(argv)
    output = build_document(args.source, args.output)
    print(f"Built {output.resolve()} at {datetime.now().isoformat(timespec='seconds')}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
