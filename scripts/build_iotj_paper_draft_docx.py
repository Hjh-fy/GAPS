"""Build the code-aligned Chinese IoT-J paper draft as a styled DOCX."""
from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = REPO_ROOT / "docs" / "paper" / "GAPS_IoTJ_paper_draft_20260711.zh.md"
DEFAULT_OUTPUT = REPO_ROOT / "docs" / "paper" / "GAPS_IoTJ_paper_draft_20260711.zh.docx"
DEFAULT_METRICS = (
    REPO_ROOT
    / "results"
    / "iotj_classification_ablation_20260711_v2r1_summary"
    / "classification_per_run.csv"
)
ASSET_DIR = REPO_ROOT / "docs" / "paper" / "assets" / "gaps_iotj_draft_20260711"
MSYH_FONT = Path("C:/Windows/Fonts/msyh.ttc")

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
HEADER_FOOTER_DXA = 708

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEXT = "202832"
MUTED = "66717D"
LIGHT_FILL = "F4F6F9"
LIGHT_BLUE = "DCEBF7"
LIGHT_GREEN = "DCEFE4"
LIGHT_CORAL = "F8E2D7"
LIGHT_GOLD = "F8EBC7"
WHITE = "FFFFFF"
BLACK = "000000"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(
    run,
    *,
    ascii_font: str = "Calibri",
    east_asia_font: str = "Microsoft YaHei",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_bottom_border(paragraph, *, color: str, size: int = 10, space: int = 5) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_paragraph_callout_box(paragraph, *, fill: str, border: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)

    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "5")
        node.set(qn("w:color"), border)
        p_bdr.append(node)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    set_run_font(run, size=8.5, color=MUTED)


def configure_header_footer(section) -> None:
    section.header_distance = Twips(HEADER_FOOTER_DXA)
    section.footer_distance = Twips(HEADER_FOOTER_DXA)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("GAPS 云边协同气体感知系统")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    p.add_run("\t")
    right = p.add_run("IoT-J 中文论文初稿")
    set_run_font(right, size=8.5, color=MUTED)
    set_paragraph_bottom_border(p, color="C9D4DE", size=6, space=4)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(3)
    prefix = fp.add_run("Updated draft 2026-07-15  |  Page ")
    set_run_font(prefix, size=8.5, color=MUTED)
    add_page_field(fp)


def set_style_font(style, *, size: float, color: str, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=TEXT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = False

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, size=13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    if "Paper Caption" not in doc.styles:
        caption = doc.styles.add_style("Paper Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Paper Caption"]
    set_style_font(caption, size=9, color=MUTED)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    if "Paper Equation" not in doc.styles:
        equation = doc.styles.add_style("Paper Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = doc.styles["Paper Equation"]
    set_style_font(equation, size=10.5, color=BLACK)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(6)
    equation.paragraph_format.space_after = Pt(6)
    equation.paragraph_format.keep_together = True

    if "Table Citation" not in doc.styles:
        table_citation = doc.styles.add_style("Table Citation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_citation = doc.styles["Table Citation"]
    set_style_font(table_citation, size=9, color=MUTED)
    table_citation.paragraph_format.space_before = Pt(4)
    table_citation.paragraph_format.space_after = Pt(4)


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        configure_header_footer(section)
    configure_styles(doc)
    doc.core_properties.title = "面向传感器漂移的校准辅助云边协同气体感知系统"
    doc.core_properties.subject = "IEEE Internet of Things Journal Chinese paper draft"
    doc.core_properties.author = "GAPS Research Team"
    doc.core_properties.keywords = "federated learning; gas sensing; sensor drift; cloud-edge; calibration"


def next_numbering_ids(numbering) -> tuple[int, int]:
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    return (max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1)


def create_numbering(doc: Document, *, kind: str, bracket: bool = False) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id, num_id = next_numbering_ids(numbering)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{0xA1000000 + abstract_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    template = OxmlElement("w:tmpl")
    template.set(qn("w:val"), f"{0xB2000000 + num_id:08X}")
    abstract.append(template)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    num_fmt = OxmlElement("w:numFmt")
    if kind == "bullet":
        num_fmt.set(qn("w:val"), "bullet")
        level_text_value = "•"
    else:
        num_fmt.set(qn("w:val"), "decimal")
        level_text_value = "[%1]" if bracket else "%1."
    level.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), level_text_value)
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_node))
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend((color, underline))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\*\*.*?\*\*|`.*?`|https?://[^\s]+)")


def add_inline_text(paragraph, text: str, *, size: float | None = None) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                ascii_font="Consolas",
                east_asia_font="Microsoft YaHei",
                size=9.5 if size is None else min(size, 9.5),
                color=DARK_BLUE,
            )
        else:
            clean = token.rstrip(".,;，。；")
            trailing = token[len(clean) :]
            add_hyperlink(paragraph, clean, clean)
            if trailing:
                run = paragraph.add_run(trailing)
                set_run_font(run, size=size)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    add_inline_text(paragraph, text)


def add_callout(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Twips(TABLE_INDENT_DXA)
    paragraph.paragraph_format.right_indent = Twips(TABLE_INDENT_DXA)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.keep_together = True
    set_paragraph_callout_box(paragraph, fill=LIGHT_FILL, border="8A98A5")
    label = paragraph.add_run("稿件状态  ")
    set_run_font(label, size=9.5, color=BLUE, bold=True)
    add_inline_text(paragraph, text, size=9.5)


def add_equation(doc: Document, equation: str) -> None:
    paragraph = doc.add_paragraph(style="Paper Equation")
    run = paragraph.add_run(equation)
    set_run_font(
        run,
        ascii_font="Cambria Math",
        east_asia_font="Microsoft YaHei",
        size=10.5,
        color=BLACK,
    )


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def table_widths(headers: Sequence[str]) -> tuple[int, ...]:
    key = tuple(headers)
    explicit: dict[tuple[str, ...], tuple[int, ...]] = {
        ("项目", "冻结设置", "作用"): (2100, 2350, 4910),
        ("模块", "结构", "主要参数"): (1800, 3000, 4560),
        ("损失项", "权重", "当前解释边界"): (2900, 1100, 5360),
        ("ID", "客户端机制", "聚合", "服务器机制", "因果问题"): (650, 2050, 1600, 2450, 2610),
        ("类别", "参数", "设置"): (2100, 3000, 4260),
        ("组别", "主要机制", "Accuracy", "Macro-F1", "NLL", "ECE"): (700, 3260, 1350, 1350, 1350, 1350),
        ("候选", "Test RMSE", "S_CC N", "S_CC RMSE", "选择结果", "证据地位"): (1700, 1200, 900, 1300, 2100, 2160),
        ("证据模块", "已完成", "尚需完成", "最终论文输出"): (1350, 2400, 3000, 2610),
        ("当前可主张", "当前不可主张"): (4680, 4680),
    }
    if key in explicit:
        return explicit[key]
    count = len(headers)
    base = CONTENT_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return tuple(widths)


def add_markdown_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    headers = list(rows[0])
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = table_widths(headers)
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(value)
        set_run_font(run, size=8.5, color=TEXT, bold=True)
    set_repeat_table_header(table.rows[0])

    for row_values in rows[1:]:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            if index >= 2 and re.fullmatch(r"[-+]?\d+(?:\.\d+)?%?", value):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_text(paragraph, value, size=8.5)
    set_table_geometry(table, widths)
    after = doc.add_paragraph(style="Table Citation")
    after.add_run("")


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.35))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption.split(".", 1)[0])
    cap = doc.add_paragraph(style="Paper Caption")
    cap.add_run(caption)


def font_properties() -> FontProperties:
    if not MSYH_FONT.is_file():
        raise FileNotFoundError(f"Chinese font not found: {MSYH_FONT}")
    return FontProperties(fname=str(MSYH_FONT))


def draw_box(ax, xy, width, height, title, detail, facecolor, font) -> None:
    box = FancyBboxPatch(
        xy,
        width,
        height,
        boxstyle="round,pad=0.012,rounding_size=0.012",
        linewidth=1.2,
        edgecolor="#607080",
        facecolor=facecolor,
    )
    ax.add_patch(box)
    ax.text(
        xy[0] + width / 2,
        xy[1] + height * 0.65,
        title,
        ha="center",
        va="center",
        fontsize=11,
        fontweight="bold",
        color="#1F2A35",
        fontproperties=font,
    )
    ax.text(
        xy[0] + width / 2,
        xy[1] + height * 0.30,
        detail,
        ha="center",
        va="center",
        fontsize=8.6,
        color="#44515E",
        fontproperties=font,
        linespacing=1.25,
    )


def arrow(ax, start, end, label, font, *, color="#566675", rad=0.0) -> None:
    patch = FancyArrowPatch(
        start,
        end,
        arrowstyle="-|>",
        mutation_scale=12,
        linewidth=1.2,
        color=color,
        connectionstyle=f"arc3,rad={rad}",
    )
    ax.add_patch(patch)
    if label:
        x = (start[0] + end[0]) / 2
        y = (start[1] + end[1]) / 2 + (0.035 if rad >= 0 else -0.035)
        ax.text(
            x,
            y,
            label,
            ha="center",
            va="center",
            fontsize=7.8,
            color="#566675",
            fontproperties=font,
            bbox=dict(facecolor="white", edgecolor="none", pad=1.2),
        )


def build_architecture_figure(path: Path) -> None:
    font = font_properties()
    fig, ax = plt.subplots(figsize=(11.5, 5.7), dpi=220)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    draw_box(ax, (0.03, 0.66), 0.20, 0.20, "C1 源客户端", "物理树莓派\n本地 CE/语义训练", "#DCEBF7", font)
    draw_box(ax, (0.03, 0.28), 0.20, 0.20, "C2 源客户端", "物理 PC\n本地 CE/语义训练", "#DCEBF7", font)
    draw_box(ax, (0.34, 0.52), 0.28, 0.30, "阿里云 ECS", "Flower 聚合\nclass-phase 语义记忆\nC5 校准辅助适配", "#F8E2D7", font)
    draw_box(ax, (0.37, 0.13), 0.22, 0.18, "C5 Calibration", "320 窗口\nDA / 回归 / QC 选择", "#DCEFE4", font)
    draw_box(ax, (0.72, 0.69), 0.24, 0.17, "冻结分类器", "pred_class / logits\n64-D features", "#E8E0F2", font)
    draw_box(ax, (0.72, 0.42), 0.24, 0.17, "C5 个性化定量", "H2.3+ / H8 / selector\n按预测类别路由", "#DCEFE4", font)
    draw_box(ax, (0.72, 0.14), 0.24, 0.17, "可靠性 QC", "accept / review / reject\n自动 ppm 或人工复核", "#F8EBC7", font)

    arrow(ax, (0.23, 0.72), (0.34, 0.64), "", font)
    arrow(ax, (0.23, 0.42), (0.34, 0.58), "", font)
    arrow(ax, (0.34, 0.76), (0.23, 0.82), "", font, rad=-0.12)
    arrow(ax, (0.34, 0.55), (0.23, 0.34), "", font, rad=0.12)
    ax.text(
        0.285,
        0.52,
        "上行: 参数/语义统计\n下行: 全局模型/原型",
        ha="center",
        va="center",
        fontsize=7.6,
        color="#566675",
        fontproperties=font,
        bbox=dict(facecolor="white", edgecolor="#D9E0E6", linewidth=0.5, pad=2.0),
    )
    arrow(ax, (0.48, 0.31), (0.48, 0.52), "目标校准", font)
    arrow(ax, (0.62, 0.67), (0.72, 0.77), "round-25", font)
    arrow(ax, (0.84, 0.69), (0.84, 0.59), "路由与特征", font)
    arrow(ax, (0.59, 0.22), (0.72, 0.49), "目标拟合/阈值", font)
    arrow(ax, (0.84, 0.42), (0.84, 0.31), "浓度与风险", font)

    ax.text(0.03, 0.94, "源域边缘训练", fontsize=10, color="#2E74B5", fontweight="bold", fontproperties=font)
    ax.text(0.34, 0.94, "云端协同与目标校准", fontsize=10, color="#A65335", fontweight="bold", fontproperties=font)
    ax.text(0.72, 0.94, "目标端推理与可靠输出", fontsize=10, color="#287A4A", fontweight="bold", fontproperties=font)
    fig.tight_layout(pad=0.5)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def build_evaluation_figure(path: Path) -> None:
    font = font_properties()
    fig, ax = plt.subplots(figsize=(11.2, 4.2), dpi=220)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    draw_box(ax, (0.04, 0.57), 0.18, 0.22, "分类指标", "Accuracy / F1\nNLL / ECE", "#DCEBF7", font)
    draw_box(ax, (0.31, 0.57), 0.22, 0.22, "S_CC 能力集", "分类正确\nQC 之前", "#E8E0F2", font)
    draw_box(ax, (0.63, 0.57), 0.30, 0.22, "回归能力线", "RMSE / MAE / NRMSE / R²\n回答数值模型上限", "#DCEFE4", font)

    draw_box(ax, (0.04, 0.16), 0.18, 0.22, "真实预测路由", "全部 C5 test\nS_ALL", "#F8E2D7", font)
    draw_box(ax, (0.31, 0.16), 0.22, 0.22, "可靠性 QC", "accept / review / reject\n阈值仅用 validation", "#F8EBC7", font)
    draw_box(ax, (0.63, 0.16), 0.30, 0.22, "系统交付线", "误差 + N + coverage\n随机拒绝/固定覆盖率对照", "#DCEFE4", font)

    arrow(ax, (0.22, 0.68), (0.31, 0.68), "条件筛选", font)
    arrow(ax, (0.53, 0.68), (0.63, 0.68), "per-gas regression", font)
    arrow(ax, (0.22, 0.27), (0.31, 0.27), "风险计算", font)
    arrow(ax, (0.53, 0.27), (0.63, 0.27), "三态输出", font)

    ax.text(0.04, 0.90, "能力线: 不用 QC 美化模型能力", fontsize=10, color="#2E74B5", fontweight="bold", fontproperties=font)
    ax.text(0.04, 0.47, "系统线: 不隐藏错路由与筛选覆盖率", fontsize=10, color="#A65335", fontweight="bold", fontproperties=font)
    fig.tight_layout(pad=0.4)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def read_classification_metrics(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))
    by_group = {row["group_id"]: row for row in rows}
    order = ["A0", "A0T", "A2", "A3", "A4", "A4S", "A5", "A6", "A7"]
    missing = [group for group in order if group not in by_group]
    if missing:
        raise ValueError(f"classification summary is missing groups: {missing}")
    return [by_group[group] for group in order]


def build_classification_figure(path: Path, metrics_path: Path) -> None:
    rows = read_classification_metrics(metrics_path)
    font = font_properties()
    groups = [row["group_id"] for row in rows]
    accuracy = [100 * float(row["accuracy"]) for row in rows]
    macro_f1 = [100 * float(row["macro_f1"]) for row in rows]
    colors = ["#7E8B97", "#C85E4A", "#6A8EB5", "#6A8EB5", "#6A8EB5", "#8C6DAA", "#D59436", "#3D9A68", "#B4464A"]

    fig, ax = plt.subplots(figsize=(10.8, 5.8), dpi=220)
    y_positions = list(range(len(groups)))
    bar_height = 0.34
    acc_bars = ax.barh(
        [y - bar_height / 2 for y in y_positions],
        accuracy,
        height=bar_height,
        color=colors,
        edgecolor="white",
        linewidth=0.6,
        label="Accuracy",
    )
    f1_bars = ax.barh(
        [y + bar_height / 2 for y in y_positions],
        macro_f1,
        height=bar_height,
        color=colors,
        alpha=0.52,
        edgecolor="#35414C",
        linewidth=0.5,
        hatch="//",
        label="Macro-F1",
    )
    ax.set_yticks(y_positions, groups, fontproperties=font, fontsize=9.5)
    ax.invert_yaxis()
    ax.set_xlim(0, 105)
    ax.set_xlabel("C5 test performance (%)", fontproperties=font, fontsize=9.5)
    ax.set_title("C1/C2 → C5 分类核心筛选 (seed 42, round 25)", fontproperties=font, fontsize=12, fontweight="bold")
    ax.xaxis.grid(True, color="#D9E0E6", linewidth=0.7)
    ax.set_axisbelow(True)
    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.spines["bottom"].set_color("#8A98A5")
    ax.tick_params(axis="x", labelsize=8.5, colors="#44515E")

    for bars, values in ((acc_bars, accuracy), (f1_bars, macro_f1)):
        for bar, value in zip(bars, values):
            x = min(value + 1.0, 101.5)
            ax.text(
                x,
                bar.get_y() + bar.get_height() / 2,
                f"{value:.2f}",
                va="center",
                ha="left",
                fontsize=7.4,
                color="#26323D",
                fontproperties=font,
            )
    ax.legend(
        loc="upper center",
        bbox_to_anchor=(0.55, 0.985),
        ncol=2,
        frameon=False,
        prop=font,
    )
    ax.text(
        0.0,
        -0.12,
        "注: A0T 使用相同 C5 标签预算的 target-CE; A6 与 A5 为替换式方法族对照; 当前未包含多种子不确定性。",
        transform=ax.transAxes,
        fontsize=7.8,
        color="#5F6B76",
        fontproperties=font,
    )
    fig.tight_layout(pad=0.8)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def generate_figures(
    metrics_path: Path,
    *,
    reuse_existing: bool = False,
) -> dict[str, tuple[Path, str]]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "system_architecture.png"
    evaluation = ASSET_DIR / "evaluation_contract.png"
    classification = ASSET_DIR / "classification_seed42.png"
    figure_paths = (architecture, evaluation, classification)
    if reuse_existing:
        missing = [path for path in figure_paths if not path.is_file()]
        if missing:
            raise FileNotFoundError(f"cannot reuse missing figure assets: {missing}")
    else:
        global plt, FontProperties, FancyArrowPatch, FancyBboxPatch
        import matplotlib.pyplot as plt
        from matplotlib.font_manager import FontProperties
        from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

        build_architecture_figure(architecture)
        build_evaluation_figure(evaluation)
        build_classification_figure(classification, metrics_path)
    return {
        "architecture": (architecture, "Fig. 1. GAPS 真实云边训练、C5 校准和可靠输出系统架构。"),
        "evaluation_contract": (evaluation, "Fig. 2. 能力线与真实路由系统线的双结果合同。"),
        "classification": (classification, "Fig. 3. 九组 seed-42 round-25 C5 分类核心筛选结果。"),
    }


def extract_metadata(lines: Sequence[str]) -> tuple[str, str, str, int]:
    title = ""
    subtitle = ""
    english_title = ""
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if line.startswith("# "):
            title = line[2:].strip()
        elif line == "## 副标题":
            subtitle = next((item.strip() for item in lines[index + 1 :] if item.strip()), "")
        elif line == "## English Title":
            english_title = next((item.strip() for item in lines[index + 1 :] if item.strip()), "")
        elif line == "> [!STATUS]":
            return title, subtitle, english_title, index
        index += 1
    raise ValueError("paper source is missing the status callout")


def add_memo_masthead(doc: Document, title: str, subtitle: str, english_title: str) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(2)
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("IEEE INTERNET OF THINGS JOURNAL  |  CHINESE COMPLETE DRAFT")
    set_run_font(run, size=9, color=BLUE, bold=True)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(5)
    title_p.paragraph_format.line_spacing = 1.08
    title_p.paragraph_format.keep_with_next = True
    display_title = title.replace("云边协同", "\n云边协同", 1)
    title_run = title_p.add_run(display_title)
    set_run_font(title_run, size=21.5, color=BLACK, bold=True)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(5)
    subtitle_p.paragraph_format.keep_with_next = True
    subtitle_run = subtitle_p.add_run(subtitle)
    set_run_font(subtitle_run, size=13, color=DARK_BLUE, bold=True)

    en_p = doc.add_paragraph()
    en_p.paragraph_format.space_before = Pt(0)
    en_p.paragraph_format.space_after = Pt(12)
    en_p.paragraph_format.keep_with_next = True
    en_run = en_p.add_run(english_title)
    set_run_font(en_run, size=10.5, color=MUTED, italic=True)

    metadata = [
        ("稿件类型", "IoT-J 系统方法论文完整初稿"),
        ("主协议", "C1/C2 source → C5 target only"),
        ("证据版本", "v2r1 seed-42 core complete; confirmation pending"),
        ("日期", "2026-07-15"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=9.5, color=TEXT, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=9.5, color=TEXT)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(3)
    rule.paragraph_format.space_after = Pt(10)
    set_paragraph_bottom_border(rule, color=BLUE, size=14, space=5)


def is_separator_row(cells: Sequence[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def parse_markdown_into_docx(
    doc: Document,
    lines: Sequence[str],
    start_index: int,
    figures: dict[str, tuple[Path, str]],
) -> None:
    index = start_index
    current_section = ""
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue

        if line == "> [!STATUS]":
            callout_lines: list[str] = []
            index += 1
            while index < len(lines) and lines[index].lstrip().startswith(">"):
                callout_lines.append(lines[index].lstrip()[1:].strip())
                index += 1
            add_callout(doc, " ".join(callout_lines))
            continue

        if line == "[[PAGEBREAK]]":
            paragraph = doc.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            index += 1
            continue

        figure_match = re.fullmatch(r"\[\[FIGURE:([a-z_]+)\]\]", line)
        if figure_match:
            key = figure_match.group(1)
            if key not in figures:
                raise KeyError(f"unknown figure key: {key}")
            add_figure(doc, *figures[key])
            index += 1
            continue

        if line.startswith("### "):
            current_section = line[4:].strip()
            doc.add_paragraph(current_section, style="Heading 2")
            index += 1
            continue

        if line.startswith("## "):
            current_section = line[3:].strip()
            doc.add_paragraph(current_section, style="Heading 1")
            index += 1
            continue

        if line.startswith("$$"):
            equation_parts = [line]
            while not equation_parts[-1].endswith("$$") or len(equation_parts[-1]) == 2:
                index += 1
                equation_parts.append(lines[index].strip())
            equation = " ".join(equation_parts).strip("$").strip()
            add_equation(doc, equation)
            index += 1
            continue

        if line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            parsed = [parse_table_row(item) for item in table_lines]
            parsed = [row for row in parsed if not is_separator_row(row)]
            add_markdown_table(doc, parsed)
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            items: list[str] = []
            while index < len(lines):
                match = re.match(r"^\d+\.\s+(.*)$", lines[index].strip())
                if not match:
                    break
                items.append(match.group(1))
                index += 1
            references = current_section == "参考文献"
            num_id = create_numbering(doc, kind="decimal", bracket=references)
            for item in items:
                paragraph = doc.add_paragraph(style="Normal")
                apply_numbering(paragraph, num_id)
                if references:
                    paragraph.paragraph_format.line_spacing = 1.05
                    paragraph.paragraph_format.space_after = Pt(3)
                    add_inline_text(paragraph, item, size=9)
                else:
                    add_inline_text(paragraph, item)
            continue

        bullet = re.match(r"^-\s+(.*)$", line)
        if bullet:
            items = []
            while index < len(lines):
                match = re.match(r"^-\s+(.*)$", lines[index].strip())
                if not match:
                    break
                items.append(match.group(1))
                index += 1
            num_id = create_numbering(doc, kind="bullet")
            for item in items:
                paragraph = doc.add_paragraph(style="Normal")
                apply_numbering(paragraph, num_id)
                add_inline_text(paragraph, item)
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line:
                break
            if (
                next_line.startswith(("## ", "### ", "|", "$$", ">", "[["))
                or re.match(r"^\d+\.\s+", next_line)
                or re.match(r"^-\s+", next_line)
            ):
                break
            paragraph_lines.append(next_line)
            index += 1
        add_body_paragraph(doc, " ".join(paragraph_lines))


def audit_document(doc: Document) -> None:
    section = doc.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.top_margin == Inches(1)
    assert section.right_margin == Inches(1)
    assert section.bottom_margin == Inches(1)
    assert section.left_margin == Inches(1)
    assert section.header_distance == Twips(HEADER_FOOTER_DXA)
    assert section.footer_distance == Twips(HEADER_FOOTER_DXA)

    normal = doc.styles["Normal"]
    assert normal.font.size == Pt(11)
    assert normal.paragraph_format.space_after == Pt(8)
    assert abs(float(normal.paragraph_format.line_spacing) - 1.333) < 0.005
    for style_name, size, before, after in (
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = doc.styles[style_name]
        assert style.font.size == Pt(size)
        assert style.paragraph_format.space_before == Pt(before)
        assert style.paragraph_format.space_after == Pt(after)

    if not doc.tables:
        raise AssertionError("paper draft must contain tables")
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        layout = tbl_pr.find(qn("w:tblLayout"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == CONTENT_WIDTH_DXA
        assert tbl_ind is not None and int(tbl_ind.get(qn("w:w"))) == TABLE_INDENT_DXA
        assert layout is not None and layout.get(qn("w:type")) == "fixed"
        grid_widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
        assert sum(grid_widths) == CONTENT_WIDTH_DXA

    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    for required in (
        "摘要",
        "1. 引言",
        "4. 校准辅助联邦语义分类",
        "8. 结果与分析",
        "10. 结论",
        "参考文献",
    ):
        if required not in body_text:
            raise AssertionError(f"missing required section: {required}")


def build_document(
    source: Path,
    output: Path,
    metrics: Path,
    *,
    reuse_existing_figures: bool = False,
) -> Path:
    if not source.is_file():
        raise FileNotFoundError(source)
    if not metrics.is_file():
        raise FileNotFoundError(metrics)
    lines = source.read_text(encoding="utf-8").splitlines()
    title, subtitle, english_title, start_index = extract_metadata(lines)
    figures = generate_figures(metrics, reuse_existing=reuse_existing_figures)

    doc = Document()
    configure_document(doc)
    add_memo_masthead(doc, title, subtitle, english_title)
    parse_markdown_into_docx(doc, lines, start_index, figures)
    audit_document(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    reloaded = Document(output)
    audit_document(reloaded)
    return output


def main(argv: Sequence[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--metrics", type=Path, default=DEFAULT_METRICS)
    parser.add_argument(
        "--reuse-existing-figures",
        action="store_true",
        help="reuse pre-rendered figure assets instead of importing matplotlib",
    )
    args = parser.parse_args(argv)
    output = build_document(
        args.source,
        args.output,
        args.metrics,
        reuse_existing_figures=args.reuse_existing_figures,
    )
    print(output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
